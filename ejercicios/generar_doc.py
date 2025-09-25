import json
import os
from docx import Document

# Crear documento
doc = Document()
doc.add_heading("📘 Documentación del Proyecto Divisas App", level=0)

# ================================
# Leer package.json
# ================================
try:
    with open("package.json", "r", encoding="utf-8") as f:
        package_data = json.load(f)
except FileNotFoundError:
    print("❌ No se encontró package.json en la carpeta actual.")
    exit()

# Info general
doc.add_heading("📌 Información del Proyecto", level=1)
doc.add_paragraph(f"Nombre: {package_data.get('name', 'N/A')}")
doc.add_paragraph(f"Versión: {package_data.get('version', 'N/A')}")
doc.add_paragraph(f"Gestor de paquetes: pnpm v. 10.13.1")
doc.add_paragraph(f"Descripción: {package_data.get('description', 'N/A')}")
# Dependencias
doc.add_heading("📦 Dependencias", level=1)
dependencies = package_data.get("dependencies", {})
dev_dependencies = package_data.get("devDependencies", {})

if dependencies:
    doc.add_paragraph("Dependencias principales:")
    for dep, ver in dependencies.items():
        doc.add_paragraph(f" - {dep}: {ver}")
else:
    doc.add_paragraph("No se encontraron dependencias principales.")

if dev_dependencies:
    doc.add_paragraph("\nDependencias de desarrollo:")
    for dep, ver in dev_dependencies.items():
        doc.add_paragraph(f" - {dep}: {ver}")
else:
    doc.add_paragraph("No se encontraron dependencias de desarrollo.")

# Scripts
doc.add_heading("🛠️ Scripts Disponibles", level=1)
scripts = package_data.get("scripts", {})
if scripts:
    for script, cmd in scripts.items():
        doc.add_paragraph(f" - npm run {script} → {cmd}")
else:
    doc.add_paragraph("No se encontraron scripts en package.json.")

# ================================
# Estructura de carpetas (excluyendo node_modules)
# ================================
doc.add_heading("📁 Estructura de Carpetas", level=1)

EXCLUIR = {"node_modules", ".git", ".vscode", "__pycache__"}

def listar_estructura(base_dir, nivel=0, max_nivel=2):
    if nivel > max_nivel:
        return ""
    estructura = ""
    try:
        for item in os.listdir(base_dir):
            if item in EXCLUIR:
                continue
            path = os.path.join(base_dir, item)
            estructura += "    " * nivel + f"- {item}\n"
            if os.path.isdir(path):
                estructura += listar_estructura(path, nivel + 1, max_nivel)
    except PermissionError:
        pass
    return estructura

estructura_proyecto = listar_estructura(".", 0, 2)
doc.add_paragraph(estructura_proyecto)

# ================================
# Explicación de la estructura
# ================================
doc.add_heading("📖 Explicación de la Estructura", level=1)

explicaciones = {
    ".env": "Variables de entorno (claves, URLs, configuraciones sensibles).",
    ".gitignore": "Archivos y carpetas que no deben subirse a Git.",
    "eslint.config.js": "Configuración de ESLint para mantener un estilo de código consistente.",
    "generar_doc.py": "Script para generar la documentación automática en Word.",
    "index.html": "Punto de entrada HTML principal para la app React.",
    "package.json": "Información del proyecto, dependencias y scripts de ejecución.",
    "pnpm-lock.yaml": "Archivo de bloqueo con versiones exactas de dependencias (PNPM).",
    "pnpm-workspace.yaml": "Configuración de workspace si se usan múltiples paquetes en el repo.",
    "public": "Archivos estáticos accesibles públicamente (íconos, fuentes, imágenes).",
    "README.md": "Documentación básica del proyecto.",
    "src": "Código fuente principal de la aplicación React.",
    "App.tsx": "Componente raíz de la aplicación.",
    "App.css": "Estilos principales del componente App.",
    "assets": "Recursos gráficos reutilizables (íconos, imágenes, fuentes).",
    "index.css": "Estilos globales del proyecto.",
    "main.tsx": "Punto de entrada de React, renderiza la app en el DOM.",
    "pages": "Páginas principales de la aplicación (login, ventas, planillas, etc.).",
    "routes": "Definición de rutas, navegación y protección de acceso.",
    "shared": "Código compartido (componentes, hooks, servicios, contextos, configuraciones).",
    "test": "Pruebas automatizadas y configuraciones de testing.",
    "vite-env.d.ts": "Declaraciones de tipos para Vite y TypeScript.",
    "tailwind.config.js": "Configuración de TailwindCSS para estilos utilitarios.",
    "tsconfig.json": "Configuración principal de TypeScript.",
    "tsconfig.app.json": "Configuración de TypeScript para la aplicación.",
    "tsconfig.node.json": "Configuración de TypeScript para Node.js.",
    "vite.config.ts": "Configuración de Vite como bundler.",
    "vitest.config.ts": "Configuración de Vitest para pruebas unitarias."
}

for archivo, descripcion in explicaciones.items():
    doc.add_paragraph(f"- {archivo} → {descripcion}")

# ================================
# Guía de ejecución
# ================================
doc.add_heading("▶️ Guía para Ejecutar el Proyecto", level=1)
doc.add_paragraph("1. Clonar el repositorio:\n   git clone <url>")
doc.add_paragraph("2. Instalar dependencias:\n   pnpm install")
doc.add_paragraph("3. Levantar el servidor de desarrollo:\n   pnpm dev")
doc.add_paragraph("4. Compilar para producción:\n   pnpm build")

# ================================
# Guardar Word
# ================================
output_file = "Documentacion_Proyecto_Divisas.docx"
doc.save(output_file)
print(f"✅ Documentación generada en '{output_file}'")